"""
Script to convert Markdown test cases to DOCX format
"""
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re
except ImportError:
    print("python-docx library not found. Installing...")
    import subprocess
    subprocess.check_call(['pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    import re

def convert_markdown_to_docx(md_file, docx_file):
    # Read markdown file
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
    
    # Create document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    lines = content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].strip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Main title (# )
        if line.startswith('# '):
            title = line[2:].strip()
            p = doc.add_heading(title, level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
        # Section headers (## )
        elif line.startswith('## '):
            header = line[3:].strip()
            doc.add_heading(header, level=1)
        
        # Subsection headers (### )
        elif line.startswith('### '):
            subheader = line[4:].strip()
            doc.add_heading(subheader, level=2)
        
        # Bold text (**text**)
        elif line.startswith('**') and ':**' in line:
            parts = line.split(':**', 1)
            label = parts[0].replace('**', '').strip()
            value = parts[1].strip() if len(parts) > 1 else ''
            
            p = doc.add_paragraph()
            run = p.add_run(label + ': ')
            run.bold = True
            p.add_run(value)
        
        # Numbered lists
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(text, style='List Number')
        
        # Horizontal rule (---)
        elif line.startswith('---'):
            doc.add_paragraph('_' * 80)
        
        # Regular paragraphs
        else:
            doc.add_paragraph(line)
        
        i += 1
    
    # Save document
    doc.save(docx_file)
    print(f"Successfully converted {md_file} to {docx_file}")

if __name__ == '__main__':
    convert_markdown_to_docx(
        'FUNCTIONAL_TEST_CASES.md',
        'FINEBANK_Functional_Test_Cases.docx'
    )
